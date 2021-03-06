
import Powerpoint_edit as powerpoint_convert
from PIL import ImageTk
from tkinter import filedialog
from tkinter import *
from docx import Document
from docx.shared import Inches
Tk().withdraw() # only draw the window needed, not the entire gui

root = Tk()
class App:

    def __init__(self, master):
        self.document = Document()

        self.numCorrect=0.0
        self.total=0.0
        self.showingReviewButton=False
        self.currentPoint=""
        self.currentPointAnswer=""
        self.reviewMode=False
        
        self.frame = Frame(master)
        
        self.frame.pack()
        self.intro_label=Label(self.frame, text="Welcome to Blanks!")
        #self.photo = PhotoImage(file="Britt_logo_0_0_1.gif")


        #self.w = Label(image=self.photo)
        #self.w.pack(side=TOP)
        self.intro_label.pack()
        #self.back_button=Button(self.frame, text="Back", command=self.back(powerpoint_convert.study_sentences_revised, powerpoint_convert.study_sentences_filled_revised))

        root.title("Britt Studios Blanks Program")
        self.restart_button=Button(self.frame, text="Restart", command=self.restart)
        self.restart_button.pack()
        self.intro_label=Label(text="Select file for conversion ")
        self.intro_label.pack()

        self.choose_file = Button(self.frame, text="Choose File", command=self.pick_file)
        self.choose_file.pack(side=LEFT)
        self.blank_intensity = Text(root, relief=RIDGE, height=1, width=1, borderwidth=3)
        self.blank_intensity.pack(side=LEFT)
        self.goButton = Button(self.frame, text="Go", command=self.go, background="cyan")
        self.goButton.config(background="blue")
        self.goButton.pack(side=LEFT)

        self.outlineButton=Button(self.frame, text="Make Outline", command=self.outline)
        self.outlineButton.pack(side=LEFT)
    def restart(self):
        self.frame.pack_forget()
        app = App(root)
        root.mainloop()
    
    def outline(self):
        for a in range(0, len(powerpoint_convert.study_sentences)-1):
            self.document.add_paragraph(powerpoint_convert.study_sentences[a])
            sentences_blanked=powerpoint_convert.study_sentences[a].split(" ")
            sentences_filled=powerpoint_convert.study_sentences_filled[a].split(" ")
            temp_study_answer=""
            for b in range(0, len(sentences_blanked)-1):
                if sentences_blanked[b] =="____":
                    temp_study_answer+=sentences_filled[b] + ", "
            p=self.document.add_paragraph("")
            p.add_run(temp_study_answer).bold=True
        outline_filename=powerpoint_convert.newFileString[0:len(powerpoint_convert.newFileString)-13]
        outline_filename+="_outline.docx"
        self.document.save(outline_filename)




    def go(self):
        self.choose_file.pack_forget()
        self.goButton.pack_forget()
        self.blank_intensity.pack_forget()
        powerpoint_convert.convert_presentation(self.filename, self.blank_intensity.get("1.0", END))
        self.study_button=Button(self.frame, text="Study", command=self.study)
        self.study_button.pack(side=LEFT)
    def reviewMissed(self):
        self.reviewMode=True
        self.study_sentences_backup=powerpoint_convert.study_sentences_revised
        self.study_sentences_filled_backup=powerpoint_convert.study_sentences_filled_revised

        self.back_button=Button(self.frame, text="Back", command=self.back)
        self.back_button.pack(side=LEFT)
        powerpoint_convert.study_sentences_revised=powerpoint_convert.missed
        powerpoint_convert.study_sentences_filled_revised=powerpoint_convert.missed_filled
        print(powerpoint_convert.study_sentences_revised)
        print(powerpoint_convert.study_sentences_filled_revised)
        #print("study_sentences_revised: ")
        #print(powerpoint_convert.study_sentences_revised)
        #print("study_sentences_filled_revised: ")
        
        #print(powerpoint_convert.study_sentences_filled_revised)
        next(self)
    def back (self):
        print("back clicked")
        self.back_button.pack_forget()
        powerpoint_convert.study_sentences_revised=self.study_sentences_backup
        powerpoint_convert.study_sentences_filled_revised=self.study_sentences_filled_backup
        #powerpoint_convert.study_blanks()
        self.study

    def study(self):
        self.study_button.pack_forget()
        if(self.total-self.numCorrect==1):
            if(self.showingReviewButton==False):
                self.review_button=Button(self.frame, text="Review Missed", command=self.reviewMissed)
                self.review_button.pack()
                self.showingReviewButton=True
       
        powerpoint_convert.study_blanks()
        

        if(len(powerpoint_convert.study_answer)==0):
            next(self)
        else:
            try:
                self.numCorrect_label['text']="Percent correct: " + (str(round(((self.numCorrect/self.total)*100), 1))+"%")
            except:
                if(self.total>0):
                    self.numCorrect_label=Label(root, text="Percent correct: " + (str(round(((self.numCorrect/self.total)*100), 1))+"%"))
                    self.numCorrect_label.pack()
                else:
                    self.numCorrect_label=Label(root, text="Percent correct: 100%")
                    self.numCorrect_label.pack()
            labelPreFormatted=powerpoint_convert.getCurrentLabel()
            labelArray=labelPreFormatted.split()
            labelFormatted=""
            totalBlanks=0
            study_answer_array=powerpoint_convert.study_sentences_filled_revised[powerpoint_convert.index].split()
            for i in range(len(labelArray)):
                if(i%8==0):
                    #print("Label Array: ")
                    #print(labelArray[i])
                    if(labelArray[i]=="____"):
                        totalBlanks=totalBlanks+1
                        #print("Total Blanks: ")
                        #print(totalBlanks)
                    if(totalBlanks<=3):
                        labelFormatted+="\n" + labelArray[i] + " "
                    else:
                        labelFormatted+="\n" + study_answer_array[i] + " "
                else:
                    #print("Label Array: ")
                    #print(labelArray[i])                    
                    if(labelArray[i]=="____"):
                        totalBlanks=totalBlanks+1
                        #print("Total Blanks: ")
                        #print(totalBlanks)
                    if(totalBlanks<=3):
                        labelFormatted+=labelArray[i] + " "
                    else:
                        labelFormatted+=study_answer_array[i] + " "

            #Create the label with the blanks
            self.currentPoint=labelFormatted
            self.currentPointAnswer=study_answer_array
            self.studyLabel=Label(root, text=labelFormatted, width=80, height=10)
            self.studyLabel.pack(side=LEFT)

            
            self.studyText=Entry(root)
            self.studyText.pack(side=LEFT)

            if(len(powerpoint_convert.getAnswer())==2):
                self.studyText2=Entry(root)
                self.studyText2.pack(side=LEFT)
            if (len(powerpoint_convert.getAnswer())>=3):
                self.studyText2=Entry(root)
                self.studyText2.pack(side=LEFT)
                self.studyText3=Entry(root)
                self.studyText3.pack(side=LEFT)

            self.studyCheck=Button(self.frame, text="Check", command=self.check)
            self.studyCheck.pack(side=LEFT)
    def check(self):
        self.studyCheck.pack_forget()
        guiAnswer=powerpoint_convert.getAnswer()
        print(guiAnswer)

        if(len(powerpoint_convert.getAnswer())==0):
            next(self)
        else:
            powerpoint_convert.study_sentences_revised.remove(powerpoint_convert.study_sentences_revised[powerpoint_convert.index])
            powerpoint_convert.study_sentences_filled_revised.remove(powerpoint_convert.study_sentences_filled_revised[powerpoint_convert.index])

        if(len(powerpoint_convert.getAnswer())==1):
            if(self.studyText.get()==guiAnswer[0]):
                    self.resultsLabel=Label(root, text="Correct")
                    #print("Correct")
                    self.numCorrect+=1.0
                    self.total+=1.0
                    self.resultsLabel.pack(side=LEFT)
            else:
                self.resultsLabel=Label(root, text="Incorrect")
                #print("Incorrect")
                self.total+=1.0
                powerpoint_convert.missed.append("".join(self.currentPoint))
                powerpoint_convert.missed_filled.append(" ".join(self.currentPointAnswer))
                #print(len(powerpoint_convert.missed))
                #print(powerpoint_convert.missed)

                #print("Missed: ")
                #print(powerpoint_convert.missed)
                #print("Missed filled: ")
                #print(powerpoint_convert.missed_filled)
                self.resultsLabel.pack(side=LEFT)
        if (len(guiAnswer)==2):
            if(self.studyText.get()==guiAnswer[0]and self.studyText2.get()==guiAnswer[1]):
                self.resultsLabel=Label(root, text="Correct")
                #print("Correct")
                self.numCorrect+=1.0
                self.total+=1.0
                self.resultsLabel.pack(side=LEFT)    
            else:
                self.resultsLabel=Label(root, text="Incorrect")
                #print("Incorrect")
                self.total+=1.0
                powerpoint_convert.missed.append("".join(self.currentPoint))
                powerpoint_convert.missed_filled.append(" ".join(self.currentPointAnswer))
                #print(len(powerpoint_convert.missed))
                #print(powerpoint_convert.missed)
                #print("Missed: ")
                #print(powerpoint_convert.missed)
                #print("Missed filled: ")
                #print(powerpoint_convert.missed_filled)
                self.resultsLabel.pack(side=LEFT)
        if (len(guiAnswer)>=3):
            if(self.studyText.get()==guiAnswer[0]and self.studyText2.get()==guiAnswer[1] and self.studyText3.get()==guiAnswer[2]):
                self.resultsLabel=Label(root, text="Correct")
                #print("Correct")
                self.numCorrect+=1.0
                self.total+=1.0

                self.resultsLabel.pack(side=LEFT)    
            else:
                self.resultsLabel=Label(root, text="Incorrect")
                #print("Incorrect")
                self.total+=1.0
                powerpoint_convert.missed.append("".join(self.currentPoint))
                powerpoint_convert.missed_filled.append(" ".join(self.currentPointAnswer))
                #print(len(powerpoint_convert.missed))
                #print(powerpoint_convert.missed)
                #print("Missed: ")
                #print(powerpoint_convert.missed)
                #print("Missed filled: ")
                #print(powerpoint_convert.missed_filled)
                self.resultsLabel.pack(side=LEFT)                  

        self.nextButton=Button(self.frame, text="Next", command=self.next)
        self.nextButton.pack(side=LEFT)

    
    def next(self):
        if(len(powerpoint_convert.study_answer)!=0):
            print("Length is not")
            self.studyText.pack_forget()
            if(len(powerpoint_convert.getAnswer())!=0):
                self.resultsLabel.pack_forget()
                self.nextButton.pack_forget()
            if(len(powerpoint_convert.getAnswer())==1):
                self.studyText.pack_forget()
            if(len(powerpoint_convert.getAnswer())==2):
                self.studyText.pack_forget()
                self.studyText2.pack_forget()
            if(len(powerpoint_convert.getAnswer())>=3):
                self.studyText.pack_forget()
                self.studyText2.pack_forget()
                self.studyText3.pack_forget()
            self.studyLabel.pack_forget()
            self.studyCheck.pack_forget()
            self.study()
        else:
            print("length is zero")
            self.study()
        
        
    def pick_file(self):
        self.filename = filedialog.askopenfilename() # show an "Open" dialog box and return the path to the selected file
        abbrreviated_name=""
        i=len(self.filename)-1
        while i>0:
            if(self.filename[i]=="/"):
                break
            else:
                abbrreviated_name=self.filename[i]+abbrreviated_name
            i-=1
            
        self.newLabel=Label(self.frame, text=abbrreviated_name)
        self.newLabel.pack(side=LEFT)
    def create_study_label(self, label):
        self.tempLabel=Label(text=label)
    def resource_path(self,relative_path):
        if hasattr(sys, '_MEIPASS'):
            return os.path.join(sys._MEIPASS, relative_path)
        return os.path.join(os.path.abspath("."), relative_path)
app = App(root)
root.mainloop()
